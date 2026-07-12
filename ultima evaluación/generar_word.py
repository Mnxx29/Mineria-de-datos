# -*- coding: utf-8 -*-
"""
Script para generar el Informe Etapa 1 - Churn en formato Word (.docx)
con formato: Arial 12, justificado, interlineado 1.0
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import re

doc = Document()

# ── Configurar márgenes ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Estilo base: Normal ──
style_normal = doc.styles['Normal']
font_normal = style_normal.font
font_normal.name = 'Arial'
font_normal.size = Pt(12)
font_normal.color.rgb = RGBColor(0, 0, 0)
pf_normal = style_normal.paragraph_format
pf_normal.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pf_normal.space_before = Pt(0)
pf_normal.space_after = Pt(6)
pf_normal.line_spacing = Pt(12)  # interlineado 1.0 para Arial 12

# ── Estilos de encabezado ──
for level, size, color in [(1, 16, RGBColor(0, 0, 0)),
                            (2, 14, RGBColor(0, 0, 0)),
                            (3, 12, RGBColor(0, 0, 0))]:
    style_name = f'Heading {level}'
    style_h = doc.styles[style_name]
    style_h.font.name = 'Arial'
    style_h.font.size = Pt(size)
    style_h.font.bold = True
    style_h.font.color.rgb = color
    style_h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_h.paragraph_format.space_before = Pt(12)
    style_h.paragraph_format.space_after = Pt(6)
    style_h.paragraph_format.line_spacing = Pt(size)
    # Quitar el estilo de fuente del tema para forzar Arial
    rpr = style_h.element.find(qn('w:rPr'))
    if rpr is not None:
        for rfonts in rpr.findall(qn('w:rFonts')):
            for attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme', 'w:cstheme']:
                if rfonts.get(qn(attr)):
                    del rfonts.attrib[qn(attr)]

# ── Estilo para viñetas ──
style_list = doc.styles.add_style('CustomBullet', WD_STYLE_TYPE.PARAGRAPH)
style_list.base_style = doc.styles['Normal']
style_list.font.name = 'Arial'
style_list.font.size = Pt(12)
style_list.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
style_list.paragraph_format.left_indent = Cm(1.27)
style_list.paragraph_format.first_line_indent = Cm(-0.63)
style_list.paragraph_format.space_before = Pt(0)
style_list.paragraph_format.space_after = Pt(4)
style_list.paragraph_format.line_spacing = Pt(12)


def add_formatted_text(paragraph, text):
    """Agrega texto con negritas procesadas desde markdown (**texto**)."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(12)
        else:
            run = paragraph.add_run(part)
            run.font.name = 'Arial'
            run.font.size = Pt(12)


# ═══════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════

# Espacios antes de la portada
for _ in range(4):
    doc.add_paragraph('', style='Normal')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Universidad de los Lagos')
r.font.name = 'Arial'
r.font.size = Pt(18)
r.bold = True

doc.add_paragraph('', style='Normal')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Minería de Datos')
r.font.name = 'Arial'
r.font.size = Pt(16)
r.bold = True

doc.add_paragraph('', style='Normal')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Evaluación Parcial 3')
r.font.name = 'Arial'
r.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Informe de Planificación y Conceptos')
r.font.name = 'Arial'
r.font.size = Pt(14)

doc.add_paragraph('', style='Normal')
doc.add_paragraph('', style='Normal')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Caso de Estudio: Predicción de Fuga de Clientes (Churn)\nen Telecomunicaciones')
r.font.name = 'Arial'
r.font.size = Pt(13)
r.italic = True

for _ in range(4):
    doc.add_paragraph('', style='Normal')

# Datos del estudiante
for line in ['Estudiante: Germán Aravena',
             'Profesor: Pablo Silva',
             'Fecha de entrega: 10 de julio de 2026']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.name = 'Arial'
    r.font.size = Pt(12)

# Salto de página
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ═══════════════════════════════════════════════════════════════

doc.add_heading('1. Introducción', level=1)

intro_paragraphs = [
    'En la industria de las telecomunicaciones, la retención de clientes constituye uno de los pilares fundamentales para la sostenibilidad financiera de las compañías. La pérdida de clientes, conocida técnicamente como Churn, genera un impacto directo en los ingresos y eleva considerablemente los costos operativos, dado que adquirir un nuevo cliente puede resultar entre cinco y siete veces más costoso que mantener a uno existente.',

    'Frente a este desafío, las técnicas de minería de datos ofrecen una oportunidad estratégica para identificar patrones de comportamiento que anticipen la decisión de abandono del servicio. El presente informe aborda el caso de una compañía de telecomunicaciones que ha recopilado datos históricos de 7.043 usuarios en el estado de California, con el propósito de desarrollar modelos predictivos que permitan al equipo de marketing actuar de manera proactiva mediante campañas de retención dirigidas.',

    'El proyecto se enmarca dentro de la metodología CRISP-DM (Cross Industry Standard Process for Data Mining), un estándar ampliamente adoptado en la industria para estructurar proyectos de ciencia de datos. Este informe corresponde a la primera etapa del proyecto y cubre las fases conceptuales y de planificación: desde la comprensión del problema de negocio hasta la propuesta de los modelos que se implementarán en la etapa práctica. Se emplearán dos enfoques complementarios: un modelo de aprendizaje supervisado basado en regresión logística para predecir la fuga, y un modelo de aprendizaje no supervisado basado en el algoritmo Apriori para descubrir reglas de asociación entre los servicios contratados y el abandono.',

    'El documento se estructura de la siguiente manera: en primer lugar se analiza el contexto del negocio y sus objetivos; luego se describen los datos disponibles; posteriormente se presenta el plan de preparación de datos, la propuesta de modelado, las hipótesis de negocio a validar y el planteamiento de las reglas de asociación; finalmente, se presentan las conclusiones y los resultados esperados del proyecto.'
]

for text in intro_paragraphs:
    p = doc.add_paragraph(style='Normal')
    add_formatted_text(p, text)

# ═══════════════════════════════════════════════════════════════
# 2. DESARROLLO
# ═══════════════════════════════════════════════════════════════

doc.add_heading('2. Desarrollo', level=1)

# ── 2.1 Comprensión del Negocio ──
doc.add_heading('2.1 Comprensión del Negocio (Business Understanding)', level=2)

doc.add_heading('Análisis del caso de estudio', level=3)

texts_21a = [
    'La compañía de telecomunicaciones objeto de este estudio enfrenta un problema recurrente en el sector: la fuga progresiva de clientes hacia la competencia. Este fenómeno, denominado Churn, no solo implica una reducción en los ingresos recurrentes, sino que también incrementa la presión sobre las áreas comerciales para captar nuevos suscriptores que compensen las bajas, elevando significativamente los costos de adquisición.',
    'El dataset disponible contiene registros de 7.043 clientes, lo cual constituye una base de datos suficientemente robusta para la construcción de modelos predictivos confiables. Los datos incluyen información demográfica, detalles sobre los servicios contratados, características del contrato y una variable indicadora de si el cliente abandonó o no el servicio en el último mes.'
]
for t in texts_21a:
    p = doc.add_paragraph(style='Normal')
    add_formatted_text(p, t)

doc.add_heading('Identificación del problema', level=3)
p = doc.add_paragraph(style='Normal')
add_formatted_text(p, 'El problema central es la incapacidad actual de la compañía para identificar de forma anticipada qué clientes tienen mayor probabilidad de abandonar el servicio. Sin esta capacidad predictiva, las campañas de retención se aplican de manera generalizada y poco eficiente, lo que desperdicia recursos en clientes que no están en riesgo y deja desatendidos a aquellos que sí lo están.')

doc.add_heading('Objetivos organizacionales', level=3)
p = doc.add_paragraph(style='Normal')
add_formatted_text(p, 'Los objetivos del negocio asociados a este proyecto son:')

obj_org = [
    'Reducir la tasa de churn mediante la identificación temprana de clientes en riesgo de abandono.',
    'Optimizar la asignación de recursos del equipo de marketing, dirigiendo las campañas de retención hacia los segmentos con mayor probabilidad de fuga.',
    'Comprender los factores que influyen en la decisión de abandono para diseñar estrategias de mejora en los servicios ofrecidos.',
    'Aumentar el valor del ciclo de vida del cliente (Customer Lifetime Value) al prolongar su permanencia.'
]
for item in obj_org:
    p = doc.add_paragraph(style='CustomBullet')
    add_formatted_text(p, '•  ' + item)

doc.add_heading('Objetivos de minería de datos', level=3)
p = doc.add_paragraph(style='Normal')
add_formatted_text(p, 'Desde la perspectiva técnica, los objetivos de minería de datos se traducen en:')

obj_md = [
    'Construir un modelo de clasificación binaria mediante regresión logística que permita predecir si un cliente abandonará el servicio (Churn = Yes) o no (Churn = No), alcanzando un desempeño medido por un Accuracy superior al 75% y un Recall alto para la clase positiva.',
    'Implementar el algoritmo Apriori para descubrir reglas de asociación que revelen combinaciones frecuentes de servicios, tipos de contrato y métodos de pago vinculadas con la fuga de clientes, con un Lift superior a 1 que valide la significancia de las reglas encontradas.'
]
for i, item in enumerate(obj_md, 1):
    p = doc.add_paragraph(style='CustomBullet')
    add_formatted_text(p, f'{i}.  {item}')

doc.add_heading('Variables clave identificadas', level=3)
p = doc.add_paragraph(style='Normal')
add_formatted_text(p, 'A partir del análisis preliminar del caso, se identifican las siguientes variables como potencialmente más influyentes en la predicción de churn: **Contract** (tipo de contrato), **tenure** (antigüedad del cliente), **MonthlyCharges** (cargo mensual), **PaymentMethod** (método de pago) e **InternetService** (tipo de servicio de internet). La variable objetivo es **Churn**.')

# ── 2.2 Comprensión de los Datos ──
doc.add_heading('2.2 Comprensión de los Datos (Data Understanding)', level=2)

doc.add_heading('Descripción general del dataset', level=3)
p = doc.add_paragraph(style='Normal')
add_formatted_text(p, 'El dataset Telco-Customer-Churn.csv contiene 7.043 registros de clientes con 21 columnas que abarcan información demográfica, servicios contratados, detalles de facturación y la variable objetivo de abandono. A continuación, se describe cada variable del conjunto de datos:')

doc.add_heading('Variables demográficas', level=3)
vars_demo = [
    '**customerID:** Identificador único de cada cliente. Esta variable es de tipo texto y no aporta información predictiva, por lo que debe ser excluida del modelado.',
    '**gender:** Género del cliente, con valores Male (masculino) y Female (femenino). Variable categórica binaria.',
    '**SeniorCitizen:** Indica si el cliente es ciudadano de la tercera edad. Se presenta como variable numérica binaria con valores 1 (sí) y 0 (no).',
    '**Partner:** Indica si el cliente tiene pareja, con valores Yes y No. Variable categórica binaria.',
    '**Dependents:** Indica si el cliente tiene personas dependientes a su cargo, con valores Yes y No. Variable categórica binaria.'
]
for v in vars_demo:
    p = doc.add_paragraph(style='CustomBullet')
    add_formatted_text(p, '•  ' + v)

doc.add_heading('Variables de servicio telefónico', level=3)
vars_tel = [
    '**PhoneService:** Indica si el cliente tiene contratado servicio telefónico (Yes/No). Variable categórica binaria.',
    '**MultipleLines:** Especifica si el cliente tiene múltiples líneas telefónicas. Presenta tres categorías: Yes, No y No phone service. Esta última categoría está condicionada por PhoneService.'
]
for v in vars_tel:
    p = doc.add_paragraph(style='CustomBullet')
    add_formatted_text(p, '•  ' + v)

doc.add_heading('Variables de servicio de internet', level=3)
vars_inet = [
    '**InternetService:** Tipo de proveedor de internet contratado. Presenta tres categorías: DSL (línea de suscriptor digital), Fiber optic (fibra óptica) y No (sin servicio). Esta variable es considerada clave por su potencial relación con el churn.',
    '**OnlineSecurity:** Servicio de seguridad en línea (Yes/No/No internet service).',
    '**OnlineBackup:** Servicio de respaldo en línea (Yes/No/No internet service).',
    '**DeviceProtection:** Protección de dispositivos (Yes/No/No internet service).',
    '**TechSupport:** Soporte técnico (Yes/No/No internet service).',
    '**StreamingTV:** Servicio de streaming de televisión (Yes/No/No internet service).',
    '**StreamingMovies:** Servicio de streaming de películas (Yes/No/No internet service).'
]
for v in vars_inet:
    p = doc.add_paragraph(style='CustomBullet')
    add_formatted_text(p, '•  ' + v)

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, 'Estas seis variables de servicios adicionales comparten la particularidad de contener la categoría "No internet service", la cual está determinada por el valor de InternetService = No.')

doc.add_heading('Variables de contrato y facturación', level=3)
vars_contract = [
    '**Contract:** Tipo de contrato del cliente. Presenta tres categorías: Month-to-month (mes a mes), One year (un año) y Two year (dos años). Se identifica como una de las variables más críticas para la predicción de churn, ya que los contratos de corto plazo implican menor compromiso.',
    '**PaperlessBilling:** Indica si el cliente tiene facturación electrónica (Yes/No). Variable categórica binaria.',
    '**PaymentMethod:** Método de pago utilizado. Presenta cuatro categorías: Electronic check (cheque electrónico), Mailed check (cheque por correo), Bank transfer (automatic) (transferencia bancaria automática) y Credit card (automatic) (tarjeta de crédito automática).',
    '**MonthlyCharges:** Monto cargado mensualmente al cliente en dólares. Variable numérica continua.',
    '**TotalCharges:** Monto total acumulado cargado al cliente. Aunque debería ser numérica, este campo suele contener valores en formato texto y espacios en blanco que requieren conversión manual a tipo numérico (float).'
]
for v in vars_contract:
    p = doc.add_paragraph(style='CustomBullet')
    add_formatted_text(p, '•  ' + v)

doc.add_heading('Variable objetivo', level=3)
p = doc.add_paragraph(style='CustomBullet')
add_formatted_text(p, '•  **Churn:** Indica si el cliente abandonó el servicio en el último mes (Yes/No). Esta es la variable objetivo del modelo supervisado. Se anticipa un desbalance en su distribución, con aproximadamente un 26% de clientes que abandonaron frente a un 74% que permanecieron.')

doc.add_heading('Observaciones preliminares sobre la calidad de los datos', level=3)
p = doc.add_paragraph(style='Normal')
add_formatted_text(p, 'Sin haber realizado aún la exploración práctica del dataset, se pueden anticipar los siguientes aspectos relevantes sobre la calidad de los datos:')

obs = [
    'La columna TotalCharges requiere una conversión de tipo de dato, ya que probablemente se encuentra almacenada como texto (string) en lugar de número (float), lo que impedirá cálculos numéricos y generará errores si no se trata previamente.',
    'Es probable que los registros con TotalCharges vacío correspondan a clientes con tenure igual a 0, es decir, clientes que acaban de ingresar y aún no han generado cargos acumulados.',
    'Las tres variables numéricas continuas (tenure, MonthlyCharges, TotalCharges) podrían presentar valores atípicos que será necesario evaluar.',
    'El desbalance esperado en la variable Churn podría sesgar el modelo supervisado si no se aplican técnicas de balanceo.'
]
for o in obs:
    p = doc.add_paragraph(style='CustomBullet')
    add_formatted_text(p, '•  ' + o)

# ── 2.3 Plan de Preparación de Datos ──
doc.add_heading('2.3 Plan de Preparación de Datos (Data Preparation)', level=2)

doc.add_heading('Limpieza de datos', level=3)
p = doc.add_paragraph(style='Normal')
add_formatted_text(p, 'El proceso de limpieza contempla las siguientes acciones:')

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, '**Eliminación de columnas no informativas:** Se eliminará la columna customerID, ya que es un identificador único sin valor predictivo y su inclusión en el modelo introduciría ruido.')

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, '**Conversión de tipos de datos:** La columna TotalCharges será convertida de formato string a tipo numérico (float64). Los valores que no puedan convertirse (espacios en blanco o cadenas vacías) serán tratados como valores nulos para su posterior imputación.')

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, '**Verificación de duplicados:** Se verificará la existencia de registros duplicados en el dataset y, en caso de encontrarse, se eliminarán conservando la primera ocurrencia.')

doc.add_heading('Estrategia para el manejo de valores faltantes (imputaciones)', level=3)
p = doc.add_paragraph(style='Normal')
add_formatted_text(p, 'Para los valores faltantes identificados, particularmente en la columna TotalCharges, se propone la siguiente estrategia de imputación:')

imputaciones = [
    '**Método principal:** Imputación mediante la mediana de la columna. Se opta por la mediana en lugar de la media aritmética porque la mediana es robusta ante la presencia de valores atípicos y distribuciones sesgadas, lo que la convierte en un estimador más representativo del valor central.',
    '**Método alternativo a evaluar:** Dado que TotalCharges tiene una relación lógica con tenure y MonthlyCharges (el cargo total es una función del tiempo de permanencia y el cargo mensual), se evaluará si la imputación calculada como tenure × MonthlyCharges genera valores coherentes con la distribución existente.',
    '**Justificación:** Se descarta la eliminación de registros con valores faltantes como primera opción, ya que en datasets de tamaño moderado (7.043 registros), cada observación aporta información valiosa para el entrenamiento de los modelos.'
]
for item in imputaciones:
    p = doc.add_paragraph(style='CustomBullet')
    add_formatted_text(p, '•  ' + item)

doc.add_heading('Tratamiento de valores atípicos (outliers)', level=3)
p = doc.add_paragraph(style='Normal')
add_formatted_text(p, 'La detección y tratamiento de valores atípicos se realizará sobre las variables numéricas continuas: tenure, MonthlyCharges y TotalCharges.')

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, '**Método de detección:**')

deteccion = [
    'Visualización mediante boxplots (diagramas de caja) para identificar gráficamente la distribución de los datos y la presencia de valores extremos.',
    'Aplicación del método del Rango Intercuartílico (IQR): se calcularán el primer cuartil (Q1) y el tercer cuartil (Q3), definiendo como valores atípicos aquellos que se encuentren por debajo de Q1 − 1.5 × IQR o por encima de Q3 + 1.5 × IQR.',
    'Complementariamente, se utilizarán gráficos de dispersión para observar la relación entre variables y detectar outliers multivariantes.'
]
for d in deteccion:
    p = doc.add_paragraph(style='CustomBullet')
    add_formatted_text(p, '•  ' + d)

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, '**Estrategia de tratamiento:**')

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, 'Antes de aplicar cualquier técnica de corrección, se evaluará si los valores atípicos representan datos legítimos del negocio. Por ejemplo, un cliente con tenure alto y TotalCharges elevado es un comportamiento esperado y no debe eliminarse. La decisión de limpieza se fundamentará estadísticamente:')

trat = [
    'Si los outliers son datos legítimos del negocio, se conservarán sin modificación.',
    'Si los outliers representan errores de registro o valores inconsistentes, se aplicará la técnica de winsorización (capping), reemplazando los valores extremos por el valor del percentil 5 o 95 según corresponda.'
]
for t in trat:
    p = doc.add_paragraph(style='CustomBullet')
    add_formatted_text(p, '•  ' + t)

# ── 2.4 Propuesta de Modelado ──
doc.add_heading('2.4 Propuesta de Modelado', level=2)

doc.add_heading('Modelo de aprendizaje supervisado: Regresión Logística', level=3)

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, '**Definición:** La regresión logística es un algoritmo de clasificación que modela la probabilidad de pertenencia a una clase mediante la función sigmoide. Para este caso, modelará la probabilidad de que un cliente abandone el servicio (Churn = Yes) en función de sus características.')

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, '**Justificación de la elección:**')

justif_rl = [
    'Es un modelo altamente interpretable: sus coeficientes permiten identificar de manera directa qué variables tienen mayor peso en la predicción y en qué dirección influyen (aumentan o reducen la probabilidad de churn).',
    'Es eficiente computacionalmente y no requiere grandes volúmenes de datos para obtener resultados confiables.',
    'Es robusto frente a la multicolinealidad moderada, lo cual es relevante considerando que algunas variables del dataset podrían estar correlacionadas (por ejemplo, tenure y TotalCharges).',
    'Es un modelo de referencia (baseline) ampliamente utilizado en la industria para problemas de clasificación binaria, lo que facilita la comparación con otros enfoques.'
]
for j in justif_rl:
    p = doc.add_paragraph(style='CustomBullet')
    add_formatted_text(p, '•  ' + j)

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, '**Configuración planificada:**')

config = [
    'División del dataset en conjunto de entrenamiento (70%) y conjunto de prueba (30%) mediante train_test_split, con semilla fija (random_state) para garantizar la reproducibilidad.',
    'Aplicación de StandardScaler sobre las variables numéricas para estandarizar su escala y evitar que variables con rangos mayores dominen el modelo.',
    'Codificación de variables categóricas mediante Label Encoding para las binarias y One-Hot Encoding para las de múltiples categorías.',
    'Manejo del desbalance de clases mediante la técnica SMOTE (Synthetic Minority Over-sampling Technique), que genera ejemplos sintéticos de la clase minoritaria (Churn = Yes) para equilibrar la distribución antes del entrenamiento. Esto es fundamental para evitar que el modelo aprenda a predecir siempre la clase mayoritaria.'
]
for c in config:
    p = doc.add_paragraph(style='CustomBullet')
    add_formatted_text(p, '•  ' + c)

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, '**Métricas de evaluación:**')

metricas_rl = [
    '**Accuracy:** proporción total de predicciones correctas.',
    '**Precision:** de todos los clientes que el modelo predijo como churn, cuántos realmente lo fueron.',
    '**Recall (Sensibilidad):** de todos los clientes que efectivamente abandonaron, cuántos fueron detectados por el modelo. Esta métrica es especialmente relevante para el negocio, ya que un Recall bajo implica clientes en riesgo que pasan desapercibidos.',
    '**F1-Score:** media armónica entre Precision y Recall, útil para evaluar el balance entre ambas métricas.',
    '**Matriz de Confusión:** tabla que muestra los verdaderos positivos, falsos positivos, verdaderos negativos y falsos negativos, permitiendo un análisis detallado del desempeño.',
    'Se compararán las métricas del conjunto de entrenamiento (Train) contra las del conjunto de prueba (Test) para diagnosticar posibles problemas de sobreajuste (Overfitting) o subajuste (Underfitting).'
]
for m in metricas_rl:
    p = doc.add_paragraph(style='CustomBullet')
    add_formatted_text(p, '•  ' + m)

doc.add_heading('Modelo de aprendizaje no supervisado: Apriori', level=3)

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, '**Definición:** El algoritmo Apriori es una técnica de aprendizaje no supervisado diseñada para descubrir reglas de asociación en conjuntos de datos transaccionales. Opera identificando conjuntos de ítems que aparecen frecuentemente juntos y luego genera reglas del tipo "Si A, entonces B".')

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, '**Justificación de la elección:**')

justif_ap = [
    'Permite descubrir patrones ocultos de co-ocurrencia entre servicios contratados, tipo de contrato, método de pago y la fuga de clientes, sin necesidad de una variable objetivo predefinida.',
    'Las reglas generadas son fácilmente interpretables por equipos no técnicos, lo que facilita la comunicación de hallazgos al área de marketing y dirección.',
    'Complementa al modelo supervisado al revelar relaciones que un modelo de clasificación podría no capturar explícitamente.'
]
for j in justif_ap:
    p = doc.add_paragraph(style='CustomBullet')
    add_formatted_text(p, '•  ' + j)

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, '**Preparación específica para Apriori:**')

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, 'A diferencia del modelo supervisado, Apriori requiere un formato de datos particular. Se realizarán las siguientes transformaciones:')

prep_apriori = [
    'Discretización de variables numéricas continuas: tenure, MonthlyCharges y TotalCharges serán transformadas en categorías ordinales. Por ejemplo, tenure se dividirá en rangos como "Corto" (0-12 meses), "Medio" (13-36 meses) y "Largo" (más de 36 meses); MonthlyCharges se clasificará como "Bajo", "Medio" y "Alto" según percentiles o puntos de corte definidos por el contexto del negocio.',
    'Aplicación de One-Hot Encoding (variables dummy) a todo el dataset para que cada columna represente una característica binaria (0 o 1), formato requerido por el algoritmo Apriori.'
]
for i, item in enumerate(prep_apriori, 1):
    p = doc.add_paragraph(style='CustomBullet')
    add_formatted_text(p, f'{i}.  {item}')

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, '**Parámetros clave a definir:**')

params = [
    '**min_support:** umbral mínimo de frecuencia para considerar un conjunto de ítems como frecuente. Se iniciará con un valor de 0.05 (5%) y se ajustará según la cantidad de reglas generadas.',
    '**min_confidence:** confianza mínima para que una regla sea considerada relevante. Se establecerá inicialmente en 0.5 (50%).',
    '**min_lift:** se filtrarán las reglas con Lift mayor a 1, lo cual indica que la asociación es más fuerte que lo esperado por azar.'
]
for param in params:
    p = doc.add_paragraph(style='CustomBullet')
    add_formatted_text(p, '•  ' + param)

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, '**Métricas de evaluación:**')

metricas_ap = [
    '**Support (Soporte):** proporción de transacciones que contienen los ítems de la regla. Indica la frecuencia de la combinación.',
    '**Confidence (Confianza):** probabilidad de que ocurra el consecuente dado el antecedente. Mide la fuerza predictiva de la regla.',
    '**Lift (Elevación):** ratio entre la confianza observada y la confianza esperada si los ítems fueran independientes. Un Lift > 1 indica una asociación positiva significativa.'
]
for m in metricas_ap:
    p = doc.add_paragraph(style='CustomBullet')
    add_formatted_text(p, '•  ' + m)

# ── 2.5 Planteamiento de Hipótesis ──
doc.add_heading('2.5 Planteamiento de Hipótesis', level=2)

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, 'A continuación, se presentan dos hipótesis de negocio que serán validadas o refutadas mediante los modelos de aprendizaje propuestos.')

doc.add_heading('Hipótesis 1 — Modelo supervisado (Regresión Logística)', level=3)

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, '**H1:** Los clientes con contrato de tipo "Month-to-month" (mes a mes) y servicio de internet "Fiber optic" (fibra óptica) presentan una probabilidad significativamente mayor de abandonar el servicio (Churn = Yes) en comparación con clientes que poseen contratos anuales o bianuales, controlando las demás variables del modelo.')

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, '**Justificación:** Los contratos mes a mes no generan un compromiso de permanencia a largo plazo, lo que reduce la barrera de salida del cliente. Por otro lado, el servicio de fibra óptica suele estar asociado a tarifas más elevadas que el servicio DSL, lo que podría generar insatisfacción si el cliente percibe que la relación calidad-precio no es adecuada. La combinación de ambos factores (bajo compromiso contractual y costo elevado) configuraría un perfil de alto riesgo de abandono. La regresión logística permitirá cuantificar el peso individual de estas variables y verificar si su efecto es estadísticamente relevante a través de sus coeficientes.')

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, '**Forma de validación:** Se analizarán los coeficientes del modelo de regresión logística para las variables Contract_Month-to-month e InternetService_Fiber optic. Si ambos coeficientes son positivos y el modelo presenta un Recall superior al 60% para la clase Churn = Yes, se considerará la hipótesis validada.')

doc.add_heading('Hipótesis 2 — Modelo no supervisado (Apriori)', level=3)

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, '**H2:** Existen reglas de asociación significativas (con Lift mayor a 1) entre la combinación de facturación electrónica (PaperlessBilling = Yes), método de pago con cheque electrónico (PaymentMethod = Electronic check) y la fuga del cliente (Churn = Yes).')

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, '**Justificación:** Los métodos de pago que no implican un proceso automático de cobro recurrente, como el cheque electrónico, requieren una acción deliberada del cliente en cada ciclo de facturación. Esto contrasta con métodos como la transferencia bancaria automática o la tarjeta de crédito automática, donde el cobro se realiza sin intervención del usuario, generando un efecto de inercia que dificulta el abandono. La facturación electrónica, al eliminar la interacción física con el servicio (facturas en papel), podría reducir adicionalmente el vínculo percibido con la empresa. El algoritmo Apriori permitirá verificar si estas características aparecen asociadas de forma frecuente con el churn y si dicha asociación supera lo esperado por azar.')

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, '**Forma de validación:** Se buscarán reglas de asociación cuyo antecedente incluya PaperlessBilling_Yes y PaymentMethod_Electronic check, y cuyo consecuente sea Churn_Yes. Si se encuentran reglas con Support superior al 5%, Confidence superior al 50% y Lift superior a 1, se considerará la hipótesis validada.')

# ── 2.6 Reglas de Asociación ──
doc.add_heading('2.6 Reglas de Asociación y Valor para el Negocio', level=2)

doc.add_heading('¿Cómo puede el modelo Apriori ayudarnos a encontrar reglas de asociación entre los datos?', level=3)

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, 'El algoritmo Apriori es una herramienta de minería de datos que permite descubrir relaciones ocultas entre variables de un conjunto de datos, expresadas en forma de reglas de asociación del tipo "Si A, entonces B". Su aplicación en el contexto de este proyecto resulta especialmente valiosa para complementar el análisis predictivo del modelo supervisado con un enfoque descriptivo y exploratorio.')

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, '**Funcionamiento del proceso:**')

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, 'El algoritmo opera en dos etapas principales. En la primera etapa, identifica los conjuntos de ítems frecuentes, es decir, las combinaciones de características (tipo de contrato, servicios contratados, método de pago, etc.) que aparecen juntas con una frecuencia superior a un umbral mínimo definido por el parámetro de soporte (support). En la segunda etapa, genera reglas de asociación a partir de estos conjuntos frecuentes, evaluando cada regla mediante las métricas de confianza (confidence) y elevación (lift).')

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, 'Por ejemplo, el algoritmo podría descubrir una regla como: "Si un cliente tiene contrato mes a mes Y pago con cheque electrónico Y no tiene soporte técnico, entonces el cliente abandonó el servicio", con una confianza del 65% y un Lift de 2.1. Esto significa que esta combinación de características está asociada con el churn con el doble de probabilidad de lo que se esperaría si las variables fueran independientes.')

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, '**Valor que aporta al negocio:**')

p = doc.add_paragraph(style='Normal')
add_formatted_text(p, 'Las reglas de asociación generadas por Apriori entregan valor al negocio de múltiples formas:')

valor_negocio = [
    '**Identificación de perfiles de riesgo compuestos:** A diferencia de un análisis univariado (variable por variable), Apriori permite identificar combinaciones específicas de características que, en conjunto, configuran un perfil de alto riesgo. Esto habilita una segmentación más precisa de los clientes.',
    '**Interpretabilidad para equipos no técnicos:** Las reglas tienen un formato intuitivo ("Si... entonces...") que puede ser comunicado directamente a los equipos de marketing, ventas y atención al cliente sin necesidad de conocimientos estadísticos avanzados.',
    '**Diseño de estrategias de retención segmentadas:** Al conocer qué combinaciones de servicios y condiciones contractuales se asocian con mayor frecuencia al abandono, la compañía puede diseñar ofertas de retención específicas. Por ejemplo, si una regla indica que los clientes con fibra óptica, contrato mes a mes y sin soporte técnico tienen alta probabilidad de irse, se podría ofrecer a ese segmento un descuento por migrar a un contrato anual o incluir soporte técnico sin costo adicional.',
    '**Validación cruzada con el modelo supervisado:** Las reglas descubiertas por Apriori pueden confirmar o complementar las variables identificadas como más influyentes por la regresión logística, otorgando mayor robustez a las conclusiones del proyecto.',
    '**Descubrimiento de patrones no evidentes:** El modelo no supervisado puede revelar asociaciones que no habrían sido consideradas a priori, como por ejemplo que cierta combinación de servicios de streaming con un tipo específico de facturación se relaciona con la permanencia del cliente.'
]
for v in valor_negocio:
    p = doc.add_paragraph(style='CustomBullet')
    add_formatted_text(p, '•  ' + v)

# ═══════════════════════════════════════════════════════════════
# 3. CONCLUSIONES
# ═══════════════════════════════════════════════════════════════

doc.add_heading('3. Conclusiones', level=1)

conclusiones = [
    'El presente informe ha establecido las bases conceptuales y metodológicas para el desarrollo del proyecto de predicción de fuga de clientes en la compañía de telecomunicaciones. A partir del análisis de la problemática de negocio y la revisión del dataset disponible, se concluye que el proyecto es viable y que las condiciones están dadas para su implementación exitosa en la etapa práctica.',

    'En relación con la viabilidad, el dataset cuenta con 7.043 registros y 21 variables que abarcan dimensiones demográficas, de servicio, contractuales y de facturación. Este volumen y diversidad de información es suficiente para construir modelos predictivos confiables y descubrir patrones de asociación significativos.',

    'Respecto a los resultados esperados, se anticipa que las variables Contract, tenure, MonthlyCharges e InternetService serán las de mayor influencia en la predicción de churn. Se espera que la regresión logística permita clasificar correctamente a los clientes en riesgo con un Recall que priorice la detección de la clase positiva (clientes que abandonan), dado que el costo de no detectar a un cliente en fuga (falso negativo) es sustancialmente mayor que el costo de intervenir innecesariamente sobre un cliente que permanecería (falso positivo). Por su parte, se espera que el algoritmo Apriori revele combinaciones de servicios y condiciones contractuales asociadas al abandono que no son evidentes mediante un análisis univariado.',

    'La combinación de un modelo supervisado (enfoque predictivo) con un modelo no supervisado (enfoque descriptivo) permite abordar el problema desde dos ángulos complementarios: por un lado, predecir quién se irá, y por otro, entender por qué se va. Esta dualidad es fundamental para que la compañía no solo reaccione ante el churn, sino que comprenda sus causas subyacentes y diseñe estrategias preventivas efectivas.',

    'El plan de preparación de datos propuesto contempla las técnicas necesarias para garantizar que los modelos reciban datos de calidad: limpieza de valores faltantes, tratamiento de outliers fundamentado en el contexto del negocio, codificación apropiada de variables categóricas y escalamiento diferenciado según el tipo de modelo. Estas decisiones metodológicas están respaldadas por las mejores prácticas en ciencia de datos y serán implementadas rigurosamente en la siguiente etapa del proyecto.',

    'Finalmente, las dos hipótesis de negocio planteadas proporcionan un marco claro para evaluar si los modelos efectivamente responden a preguntas relevantes para la organización, asegurando que el proyecto trascienda el ejercicio técnico y se traduzca en valor estratégico para la toma de decisiones.'
]

for text in conclusiones:
    p = doc.add_paragraph(style='Normal')
    add_formatted_text(p, text)


# ── Guardar el documento ──
output_path = r'c:\Users\germa\OneDrive\Escritorio\Ing Informatica\Tercer Semestre\Mineria-de-datos\ultima evaluación\Informe_Etapa1_Churn.docx'
doc.save(output_path)
print(f'Documento generado exitosamente: {output_path}')
